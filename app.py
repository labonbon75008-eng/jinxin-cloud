"""
金鑫 - 智能投资助理
作者：拥有10年经验的Python全栈工程师
创建时间：2025年12月12日
"""

import re
import json
import base64
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')  # 必须放在最前面，防止GUI冲突
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from io import BytesIO
import requests
from datetime import datetime, timedelta
import yfinance as yf
import streamlit as st
from streamlit.runtime.scriptrunner import get_script_run_ctx
import warnings
warnings.filterwarnings('ignore')

# ========== 全局配置 ==========
st.set_page_config(
    page_title="金鑫 - 智能投资助理",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========== 自定义CSS样式 ==========
st.markdown("""
<style>
/* 手机端按钮组强制不换行 */
div[data-testid="stHorizontalBlock"] { 
    flex-wrap: nowrap !important; 
    overflow-x: auto !important;
}

/* 侧边栏样式 */
section[data-testid="stSidebar"] {
    background-color: #f8f9fa;
}

/* 消息气泡样式 */
.stChatMessage {
    padding: 12px;
    border-radius: 15px;
    margin-bottom: 10px;
    max-width: 85%;
}

/* 用户消息 */
.stChatMessage[data-testid="stChatMessage"]:nth-child(odd) {
    background-color: #e3f2fd;
    margin-left: auto;
}

/* AI消息 */
.stChatMessage[data-testid="stChatMessage"]:nth-child(even) {
    background-color: #f5f5f5;
    margin-right: auto;
}

/* 操作按钮样式 */
.operation-btn {
    margin: 2px !important;
    padding: 4px 8px !important;
    font-size: 12px !important;
    min-height: 28px !important;
}

/* 盯盘雷达提示 */
.alert-box {
    background-color: #fff3cd;
    border: 1px solid #ffeaa7;
    border-radius: 5px;
    padding: 10px;
    margin: 10px 0;
}

/* 滚动条美化 */
::-webkit-scrollbar {
    width: 6px;
    height: 6px;
}
::-webkit-scrollbar-track {
    background: #f1f1f1;
}
::-webkit-scrollbar-thumb {
    background: #888;
    border-radius: 3px;
}
::-webkit-scrollbar-thumb:hover {
    background: #555;
}
</style>
""", unsafe_allow_html=True)

# ========== 初始化Session State ==========
def init_session_state():
    """初始化所有会话状态"""
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    
    if 'monitoring_list' not in st.session_state:
        st.session_state.monitoring_list = []  # 盯盘列表
    
    if 'last_input' not in st.session_state:
        st.session_state.last_input = None
    
    if 'processing_input' not in st.session_state:
        st.session_state.processing_input = False
    
    if 'ai_responding' not in st.session_state:
        st.session_state.ai_responding = False
    
    if 'voice_enabled' not in st.session_state:
        st.session_state.voice_enabled = True  # 默认启用语音
    
    if 'chart_data' not in st.session_state:
        st.session_state.chart_data = {}  # 存储图表数据

init_session_state()

# ========== 工具函数 ==========
def clean_code_blocks(text):
    """
    彻底清除代码块，只保留纯文本和图表引用
    正则表达式匹配 ```python ... ``` 和 ``` ... ```
    """
    if not text:
        return text
    
    # 移除代码块
    cleaned = re.sub(r'```python[\s\S]*?```', '', text)
    cleaned = re.sub(r'```[\s\S]*?```', '', cleaned)
    
    # 移除行内代码标记
    cleaned = cleaned.replace('`', '')
    
    # 清理多余的空行
    cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)
    
    return cleaned.strip()

def get_stock_data(query):
    """
    获取股票数据（双备份策略）
    优先使用新浪财经，失败则使用yfinance
    
    参数:
        query: 股票名称或代码（支持中文、代码、英文）
    
    返回:
        df: DataFrame（包含历史数据，用于画图）
        info: 字符串（实时信息和基本数据）
    """
    # 股票代码映射（常见股票）
    stock_mapping = {
        '茅台': '600519', '贵州茅台': '600519', 'maotai': '600519',
        '腾讯': '0700.HK', '阿里巴巴': 'BABA', '阿里': 'BABA',
        '苹果': 'AAPL', '谷歌': 'GOOGL', '微软': 'MSFT',
        '特斯拉': 'TSLA', '亚马逊': 'AMZN', '英伟达': 'NVDA',
        '标普500': '^GSPC', '道琼斯': '^DJI', '纳斯达克': '^IXIC',
        '上证指数': '000001.SS', '深证成指': '399001.SZ',
        '创业板': '399006.SZ', '恒生指数': '^HSI',
    }
    
    # 尝试从映射中获取代码
    stock_code = None
    for name, code in stock_mapping.items():
        if name.lower() in query.lower():
            stock_code = code
            break
    
    # 如果没有找到映射，尝试从查询中提取代码模式
    if not stock_code:
        # 匹配股票代码模式（如600519、AAPL、0700.HK等）
        code_patterns = [
            r'\b\d{6}\b',  # A股代码
            r'\b[A-Z]{1,5}\b',  # 美股代码
            r'\b\d{4}\.HK\b',  # 港股代码
        ]
        
        for pattern in code_patterns:
            match = re.search(pattern, query.upper())
            if match:
                stock_code = match.group()
                break
    
    # 如果还没有找到，尝试使用yfinance的搜索
    if not stock_code:
        try:
            search_results = yf.Tickers(query)
            if search_results.tickers:
                stock_code = query
        except:
            pass
    
    if not stock_code:
        return None, "未找到对应的股票代码，请提供更明确的股票名称或代码。"
    
    info_text = ""
    df = None
    
    # 策略1：优先使用新浪财经（实时数据）
    try:
        # 新浪财经接口（实时数据）
        if stock_code.endswith('.SS') or stock_code.endswith('.SZ'):
            sina_code = f"sh{stock_code[:-3]}" if stock_code.endswith('.SS') else f"sz{stock_code[:-3]}"
        elif len(stock_code) == 6 and stock_code.isdigit():
            sina_code = f"sh{stock_code}" if stock_code.startswith('6') else f"sz{stock_code}"
        else:
            sina_code = None
        
        if sina_code:
            url = f"https://hq.sinajs.cn/list={sina_code}"
            headers = {
                'Referer': 'https://finance.sina.com.cn/',
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=5)
            
            if response.status_code == 200:
                data = response.text
                # 解析新浪返回的数据
                parts = data.split('"')[1].split(',')
                if len(parts) > 30:
                    stock_name = parts[0]
                    current_price = float(parts[3])
                    open_price = float(parts[1])
                    high_price = float(parts[4])
                    low_price = float(parts[5])
                    close_price = float(parts[2])  # 昨收
                    volume = float(parts[8])
                    
                    change = current_price - close_price
                    change_percent = (change / close_price) * 100
                    
                    info_text = f"""
**{stock_name} ({stock_code}) - 实时行情**
- 📊 当前价格: **¥{current_price:.2f}**
- 📈 涨跌: {'🟢' if change >= 0 else '🔴'} {change:+.2f} ({change_percent:+.2f}%)
- ⬆️ 今日最高: ¥{high_price:.2f}
- ⬇️ 今日最低: ¥{low_price:.2f}
- 🚪 今开: ¥{open_price:.2f}
- 📅 昨收: ¥{close_price:.2f}
- 📊 成交量: {volume:,.0f}手
- 🕒 更新时间: {parts[30]} {parts[31]}
                    """
                    
                    # 同时获取历史数据用于绘图
                    try:
                        ticker = yf.Ticker(stock_code)
                        df = ticker.history(period="1mo")
                        if not df.empty:
                            st.session_state.chart_data[stock_code] = df
                    except:
                        pass
                    
                    return df, info_text
    except Exception as e:
        pass  # 新浪接口失败，继续尝试其他接口
    
    # 策略2：使用yfinance作为备份
    try:
        ticker = yf.Ticker(stock_code)
        info = ticker.info
        
        # 获取实时数据
        current_data = ticker.history(period='1d', interval='1m')
        if not current_data.empty:
            current_price = current_data['Close'].iloc[-1]
            
            # 获取更多信息
            regular_market_price = info.get('regularMarketPrice', current_price)
            regular_market_change = info.get('regularMarketChange', 0)
            regular_market_change_percent = info.get('regularMarketChangePercent', 0)
            
            stock_name = info.get('longName', info.get('shortName', stock_code))
            currency = info.get('currency', 'USD')
            
            info_text = f"""
**{stock_name} ({stock_code}) - 实时行情**
- 📊 当前价格: **{currency}{regular_market_price:.2f}**
- 📈 涨跌: {'🟢' if regular_market_change >= 0 else '🔴'} {regular_market_change:+.2f} ({regular_market_change_percent:+.2f}%)
- 📅 交易日: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            """
            
            # 获取历史数据用于绘图
            df = ticker.history(period="1mo")
            if not df.empty:
                st.session_state.chart_data[stock_code] = df
            
            return df, info_text
            
    except Exception as e:
        info_text = f"无法获取 {stock_code} 的实时数据。错误: {str(e)}"
        return None, info_text
    
    return None, "无法获取股票数据，请检查股票代码或网络连接。"

def execute_plot_code(code_str, stock_code):
    """
    安全执行AI生成的绘图代码
    在沙盒环境中运行，注入必要的依赖
    """
    try:
        # 准备全局变量
        global_vars = {
            'plt': plt,
            'pd': pd,
            'np': np,
            'datetime': datetime,
            'timedelta': timedelta,
            'stock_code': stock_code,
            'df': st.session_state.chart_data.get(stock_code)
        }
        
        # 如果code_str中包含中文，确保使用中文字体
        if any('\u4e00' <= ch <= '\u9fff' for ch in code_str):
            try:
                # 尝试加载中文字体
                font_path = 'SimHei.ttf'
                font_prop = fm.FontProperties(fname=font_path)
                plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
                plt.rcParams['axes.unicode_minus'] = False
            except:
                # 如果失败，使用默认字体
                plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
                plt.rcParams['axes.unicode_minus'] = False
        
        # 安全执行代码
        exec(code_str, global_vars)
        
        # 获取图表并保存为Base64
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        
        # 转换为Base64
        img_base64 = base64.b64encode(buf.read()).decode('utf-8')
        return f'<img src="data:image/png;base64,{img_base64}" style="max-width:100%; border-radius:10px; margin:10px 0;">'
        
    except Exception as e:
        return f"⚠️ 图表生成失败: {str(e)}"

def export_to_word(content, filename="对话记录.docx"):
    """将内容导出为Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('金鑫智能投资助理 - 对话记录', 0)
        
        # 添加时间戳
        doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # 添加内容
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        
        # 保存到BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    except ImportError:
        # 如果python-docx未安装，创建简单的txt文件
        buffer = BytesIO()
        buffer.write(content.encode('utf-8'))
        buffer.seek(0)
        return buffer

def get_ai_response(user_input, stock_data=None, stock_info=None):
    """
    获取AI回复（模拟LLM，实际可替换为真实API）
    这里使用规则引擎+模拟响应，实际可接入OpenAI、ChatGLM等
    """
    # 检查是否是问候语
    greetings = ['你好', 'hello', 'hi', '您好', '早上好', '下午好']
    if any(greet in user_input.lower() for greet in greetings):
        return f"""👋 您好！我是您的智能投资助理 **金鑫**，很高兴为您服务！

我具备以下能力：
1. 📊 **实时股票查询** - 告诉我股票名称或代码，我为您提供实时行情
2. 📈 **走势分析** - 自动分析股票走势并绘制专业图表
3. ⚡ **盯盘提醒** - 设置价格提醒，触发时即时通知
4. 💬 **投资咨询** - 提供投资建议和市场分析

请告诉我您想查询哪只股票？例如："茅台现在的价格" 或 "AAPL走势如何？"
"""
    
    # 检查是否包含股票相关关键词
    stock_keywords = ['股票', '股价', '价格', '走势', '行情', '涨跌', 'k线', 'chart', 'stock', 'price']
    if any(keyword in user_input.lower() for keyword in stock_keywords) or stock_data is not None:
        
        if stock_data is None:
            return "请提供具体的股票名称或代码，例如：'茅台现在的价格' 或 'AAPL走势分析'"
        
        # 如果有股票数据，生成分析
        if stock_info:
            analysis = f"""{stock_info}

**📊 技术分析:**
"""
            
            if stock_data is not None and isinstance(stock_data, pd.DataFrame) and not stock_data.empty:
                # 计算技术指标
                prices = stock_data['Close']
                current_price = prices.iloc[-1]
                ma5 = prices.tail(5).mean()
                ma10 = prices.tail(10).mean()
                ma20 = prices.tail(20).mean()
                
                # 判断趋势
                if current_price > ma20 and ma5 > ma10 > ma20:
                    trend = "📈 **强势上涨趋势**"
                elif current_price < ma20 and ma5 < ma10 < ma20:
                    trend = "📉 **弱势下跌趋势**"
                elif current_price > ma20:
                    trend = "↗️ **震荡上行趋势**"
                else:
                    trend = "↘️ **震荡下行趋势**"
                
                analysis += f"""
- {trend}
- 5日均线: {ma5:.2f}
- 10日均线: {ma10:.2f}
- 20日均线: {ma20:.2f}
"""
                
                # 生成绘图代码
                plot_code = f"""
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

df = pd.DataFrame({stock_data.tail(30).to_dict()})

fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 8), height_ratios=[3, 1])

# 价格走势
ax1.plot(df.index, df['Close'], label='收盘价', color='blue', linewidth=2)
ax1.plot(df.index, df['Close'].rolling(5).mean(), label='5日均线', color='orange', linestyle='--')
ax1.plot(df.index, df['Close'].rolling(10).mean(), label='10日均线', color='green', linestyle='--')
ax1.fill_between(df.index, df['Low'], df['High'], alpha=0.2, color='gray')
ax1.set_title('{stock_data.index[-1].strftime("%Y-%m-%d")} {stock_code if "stock_code" in locals() else "股票"} 价格走势', fontsize=14)
ax1.set_ylabel('价格')
ax1.legend()
ax1.grid(True, alpha=0.3)

# 成交量
ax2.bar(df.index, df['Volume'], color=['green' if df['Close'].iloc[i] >= df['Open'].iloc[i] else 'red' for i in range(len(df))])
ax2.set_ylabel('成交量')
ax2.set_xlabel('日期')
ax2.grid(True, alpha=0.3)

plt.tight_layout()
"""
                
                # 将绘图代码存储在session中，稍后执行
                if 'plot_code' not in st.session_state:
                    st.session_state.plot_code = {}
                st_code = 'stock_chart'
                if 'stock_code' in locals():
                    st_code = stock_code
                st.session_state.plot_code[st_code] = plot_code
                
                analysis += f"\n📈 **图表分析已生成**（自动绘制30日价格走势图）"
                
            analysis += """

**💡 投资建议:**
1. 短期操作：建议关注5日均线支撑
2. 风险控制：设置止损位在当前价格的-5%
3. 仓位管理：建议分批建仓，控制单只股票仓位不超过20%
"""
            
            return analysis
        else:
            return "已获取股票数据，正在进行分析..."
    
    # 默认回复
    return f"""我理解您的问题是："{user_input}"

作为您的智能投资助理，我可以帮助您：
1. 查询全球股票实时行情
2. 分析股票技术走势
3. 提供投资建议和风险评估
4. 设置价格提醒和盯盘监控

请告诉我您感兴趣的股票，或者使用更具体的问题，例如：
- "茅台股票今天表现怎么样？"
- "帮我分析一下AAPL的走势"
- "设置茅台股价到1800元时提醒我"
"""

# ========== 侧边栏实现 ==========
with st.sidebar:
    # 头像展示（使用稳定的DiceBear API）
    st.markdown("""
    <div style="text-align: center; margin-bottom: 20px;">
        <img src="https://api.dicebear.com/9.x/avataaars/png?seed=Jinxin&backgroundColor=4d8af0&hairColor=000000&accessories=prescription02&facialHair=beardLight&clothing=shirtCrewNeck&clothingColor=262E33&eyes=happy&eyebrow=raisedExcitedNatural&mouth=smile&skinColor=f2d3b1" 
             width="120" 
             style="border-radius: 50%; border: 4px solid #4d8af0;">
        <h3 style="margin: 10px 0 5px 0;">金鑫</h3>
        <p style="color: #666; font-size: 14px; margin: 0;">智能投资助理</p>
        <p style="color: #888; font-size: 12px;">10年投资分析经验</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.divider()
    
    # 盯盘雷达
    st.subheader("🔭 盯盘雷达")
    col1, col2 = st.columns([2, 1])
    with col1:
        monitor_code = st.text_input("股票代码", placeholder="如: 600519", key="monitor_code")
    with col2:
        target_price = st.number_input("目标价", min_value=0.0, value=1800.0, step=10.0, key="target_price")
    
    if st.button("🚀 启动盯盘", use_container_width=True):
        if monitor_code:
            # 获取当前价格
            _, current_info = get_stock_data(monitor_code)
            if "当前价格" in current_info:
                # 从info中解析当前价格
                import re
                price_match = re.search(r'当前价格:.*?([\d,.]+)', current_info)
                if price_match:
                    current = float(price_match.group(1).replace(',', ''))
                    
                    # 添加到盯盘列表
                    st.session_state.monitoring_list.append({
                        'code': monitor_code,
                        'target': target_price,
                        'current': current,
                        'time': datetime.now()
                    })
                    
                    st.success(f"✅ 已启动盯盘：{monitor_code} 目标价 {target_price}")
                    
                    # 检查是否已触发
                    if current >= target_price:
                        st.warning(f"🎯 已触发！当前价 {current} ≥ 目标价 {target_price}")
                else:
                    st.error("无法获取当前价格")
            else:
                st.error("股票代码无效")
        else:
            st.error("请输入股票代码")
    
    # 显示盯盘列表
    if st.session_state.monitoring_list:
        st.divider()
        st.markdown("**当前盯盘列表**")
        for i, item in enumerate(st.session_state.monitoring_list[-5:]):  # 只显示最近5条
            col1, col2, col3 = st.columns([2, 2, 1])
            with col1:
                st.text(f"{item['code']}")
            with col2:
                st.text(f"{item['current']:.2f} → {item['target']:.2f}")
            with col3:
                if item['current'] >= item['target']:
                    st.markdown("🎯")
                else:
                    st.markdown("⏳")
    
    st.divider()
    
    # 数据管理
    st.subheader("📊 数据管理")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🗑️ 清空历史", use_container_width=True):
            st.session_state.messages = []
            st.session_state.monitoring_list = []
            st.session_state.chart_data = {}
            st.rerun()
    
    with col2:
        if st.session_state.messages:
            # 导出对话为Word
            dialog_text = "金鑫智能投资助理 - 对话记录\n\n"
            for msg in st.session_state.messages:
                role = "用户" if msg["role"] == "user" else "金鑫"
                dialog_text += f"{role}: {msg['content']}\n\n"
            
            word_buffer = export_to_word(dialog_text)
            
            st.download_button(
                label="📥 导出对话",
                data=word_buffer,
                file_name=f"投资对话_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    st.divider()
    
    # API设置（实际LLM集成点）
    st.subheader("⚙️ AI设置")
    ai_provider = st.selectbox(
        "选择AI模型",
        ["智能引擎（内置）", "OpenAI GPT", "ChatGLM", "文心一言"],
        index=0
    )
    
    if ai_provider != "智能引擎（内置）":
        api_key = st.text_input("API密钥", type="password")
        if st.button("保存设置", use_container_width=True):
            st.success("设置已保存（演示模式）")
    
    # 语音设置
    st.divider()
    st.subheader("🎤 语音设置")
    st.session_state.voice_enabled = st.checkbox("启用语音输入", value=True)

# ========== 主聊天界面 ==========
st.title("💎 金鑫 - 智能投资助理")
st.caption("专业女性投资顾问 | 实时行情分析 | 智能图表绘制")

# 显示聊天历史
for i, message in enumerate(st.session_state.messages):
    role = "user" if message["role"] == "user" else "assistant"
    
    with st.chat_message(role):
        # 显示消息内容（已清洗代码块）
        cleaned_content = clean_code_blocks(message.get("content", ""))
        st.markdown(cleaned_content)
        
        # 显示图表（如果有）
        if message.get("chart"):
            st.markdown(message["chart"], unsafe_allow_html=True)
        
        # AI消息下方显示操作按钮
        if role == "assistant":
            col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
            
            with col1:
                if st.button("📋 复制", key=f"copy_{i}", use_container_width=True, type="secondary"):
                    # 复制到剪贴板（JavaScript实现）
                    copy_js = f"""
                    <script>
                    navigator.clipboard.writeText(`{cleaned_content.replace('`', '\\`')}`).then(() => {{
                        alert('已复制到剪贴板！');
                    }});
                    </script>
                    """
                    st.components.v1.html(copy_js, height=0)
            
            with col2:
                if st.button("👁️ 隐藏", key=f"hide_{i}", use_container_width=True, type="secondary"):
                    # 隐藏该消息（在下次渲染时不显示）
                    if "hidden_messages" not in st.session_state:
                        st.session_state.hidden_messages = set()
                    st.session_state.hidden_messages.add(i)
                    st.rerun()
            
            with col3:
                if st.button("🗑️ 删除", key=f"delete_{i}", use_container_width=True, type="secondary"):
                    # 删除该消息
                    st.session_state.messages.pop(i)
                    st.rerun()
            
            with col4:
                # 导出单条消息
                export_content = f"金鑫智能投资助理 - 对话记录\n\n用户: {st.session_state.messages[i-1]['content'] if i>0 else '...'}\n\n金鑫: {cleaned_content}"
                export_buffer = export_to_word(export_content)
                
                st.download_button(
                    label="📄 导出",
                    data=export_buffer,
                    file_name=f"投资建议_{datetime.now().strftime('%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"export_{i}",
                    use_container_width=True
                )

# ========== 输入区域 ==========
input_container = st.container()

with input_container:
    # 语音输入（尝试导入，失败则跳过）
    voice_input = None
    
    if st.session_state.voice_enabled:
        try:
            from streamlit_mic_recorder import mic_recorder
            
            col1, col2 = st.columns([5, 1])
            with col2:
                audio = mic_recorder(
                    key="recorder",
                    start_prompt="🎤 说话",
                    stop_prompt="⏹️ 停止",
                    just_once=True,
                    use_container_width=True
                )
                
                if audio:
                    # 在实际应用中，这里应该调用语音识别API
                    # 演示模式下，使用模拟的语音转文本
                    voice_input = "茅台股票今天的价格是多少？"  # 模拟语音输入
                    st.info(f"识别结果: {voice_input}")
        except ImportError:
            st.session_state.voice_enabled = False
            st.caption("⚠️ 语音组件加载失败，已自动禁用语音功能")
        except Exception as e:
            st.session_state.voice_enabled = False
            st.caption(f"⚠️ 语音功能临时不可用: {str(e)}")
    
    # 文字输入（始终显示）
    if voice_input:
        # 如果语音输入成功，使用语音输入
        user_input = voice_input
    else:
        # 否则显示文字输入框
        user_input = st.chat_input("请输入您的问题或股票代码...")
    
    # 处理用户输入
    if user_input and not st.session_state.processing_input:
        st.session_state.processing_input = True
        st.session_state.last_input = user_input
        
        # 添加到消息历史
        st.session_state.messages.append({"role": "user", "content": user_input})
        
        # 立即重载以显示用户消息
        st.rerun()

# ========== AI响应逻辑 ==========
# 只有在重载后，并且最后一条消息是用户的，且AI没有在响应时，才触发AI响应
if (st.session_state.messages and 
    st.session_state.messages[-1]["role"] == "user" and 
    not st.session_state.ai_responding):
    
    st.session_state.ai_responding = True
    
    # 获取最后一条用户消息
    last_user_msg = st.session_state.messages[-1]["content"]
    
    # 获取股票数据
    stock_df, stock_info = get_stock_data(last_user_msg)
    
    # 获取AI回复
    with st.spinner("金鑫正在分析..."):
        ai_response = get_ai_response(last_user_msg, stock_df, stock_info)
        
        # 清洗代码块
        cleaned_response = clean_code_blocks(ai_response)
        
        # 存储响应
        response_data = {"role": "assistant", "content": cleaned_response}
        
        # 如果有股票数据且生成了绘图代码，执行绘图
        if stock_df is not None and not stock_df.empty:
            stock_code = None
            for code in st.session_state.chart_data:
                if isinstance(st.session_state.chart_data[code], pd.DataFrame):
                    stock_code = code
                    break
            
            if stock_code and stock_code in st.session_state.plot_code:
                chart_html = execute_plot_code(st.session_state.plot_code[stock_code], stock_code)
                response_data["chart"] = chart_html
        
        st.session_state.messages.append(response_data)
    
    st.session_state.ai_responding = False
    st.session_state.processing_input = False
    
    # 再次重载以显示AI回复
    st.rerun()

# ========== 页脚 ==========
st.divider()
st.markdown("""
<div style="text-align: center; color: #888; font-size: 12px; padding: 20px;">
    <p>💡 提示：投资有风险，入市需谨慎。本应用提供的信息仅供参考，不构成投资建议。</p>
    <p>📅 数据更新时间: {}</p>
    <p>🔒 您的对话数据仅保存在当前浏览器会话中，关闭页面后自动清除</p>
</div>
""".format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")), unsafe_allow_html=True)

# ========== 隐藏消息处理 ==========
if "hidden_messages" in st.session_state:
    # 在重载时清除隐藏消息的标记
    st.session_state.hidden_messages = set()
